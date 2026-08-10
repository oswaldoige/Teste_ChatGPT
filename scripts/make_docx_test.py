from __future__ import annotations

import json
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def resolve_project_root() -> Path:
    env_root = os.getenv("AUTOMACAO_ROOT")
    if env_root:
        return Path(env_root).expanduser().resolve()

    current = Path(__file__).resolve()
    for parent in (current.parent, *current.parents):
        if (parent / "02_EXECUCAO_OPERACIONAL").exists() and (parent / "modelos").exists():
            return parent

    return current.parent.parent


AUTOMACAO_ROOT = resolve_project_root()
STRICT_OUTPUT_DIRS = (
    AUTOMACAO_ROOT / "saida_docx",
    AUTOMACAO_ROOT / "02_EXECUCAO_OPERACIONAL" / "saida_docx",
)
STRICT_GATE_KEYS = (
    "sistema",
    "fase_e_ato_cabivel",
    "ultimo_ato_relevante",
    "peticoes_posteriores",
    "pendencia_condicionante",
    "modelo_base_path",
)


def set_run_font(run, font_name: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def apply_normal_style(document: Document, font_name: str = "Times New Roman", size_pt: int = 12) -> None:
    style = document.styles["Normal"]
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size_pt)


def find_suspicious_encoding_markers(text: str) -> list[str]:
    issues: list[str] = []
    if "\ufffd" in text:
        issues.append("caractere de substituicao Unicode (�)")

    # Captura mojibake comum em palavras, como A??o, R?PLICA, N?cleo, ?gua, 3?o.
    for match in re.finditer(r"\w\?+\w", text):
        snippet = text[max(0, match.start() - 12) : min(len(text), match.end() + 12)]
        issues.append(f"sequencia suspeita dentro de palavra: {snippet!r}")
        if len(issues) >= 5:
            break

    return issues


def validate_ptbr_text(paragraphs: list[str]) -> None:
    joined = "\n".join(paragraphs)
    issues = find_suspicious_encoding_markers(joined)
    if issues:
        details = "; ".join(issues)
        raise ValueError(
            "Texto com possivel corrompimento de acentuacao/pontuacao. "
            "Regenerar o conteudo em UTF-8 antes de salvar o .docx. "
            f"Marcadores: {details}"
        )


def normalize_abs(path: Path) -> str:
    return os.path.normcase(os.path.abspath(str(path)))


def is_path_within(path: Path, parent: Path) -> bool:
    path_norm = normalize_abs(path)
    parent_norm = normalize_abs(parent)
    return path_norm == parent_norm or path_norm.startswith(parent_norm + os.sep)


def output_requires_strict_gate(output_path: Path) -> bool:
    return any(is_path_within(output_path, strict_dir) for strict_dir in STRICT_OUTPUT_DIRS)


def derive_gate_path(input_path: Path) -> Path:
    return input_path.with_suffix(".gate.json")


def require_non_empty_string(data: dict[str, object], key: str) -> str:
    value = data.get(key)
    if not isinstance(value, str) or not value.strip():
        raise ValueError(f"Campo obrigatorio do gate ausente ou vazio: {key}")
    return value.strip()


def require_non_empty_peticoes(data: dict[str, object]) -> None:
    value = data.get("peticoes_posteriores")
    if isinstance(value, str):
        if not value.strip():
            raise ValueError("Campo obrigatorio do gate ausente ou vazio: peticoes_posteriores")
        return
    if isinstance(value, list):
        entries = [str(item).strip() for item in value if str(item).strip()]
        if not entries:
            raise ValueError("Campo obrigatorio do gate ausente ou vazio: peticoes_posteriores")
        return
    raise ValueError("Campo obrigatorio do gate em formato invalido: peticoes_posteriores")


def validate_model_base_path(model_path_str: str) -> None:
    model_path = Path(model_path_str)
    if not model_path.exists():
        raise ValueError(f"Caminho do modelo-base informado no gate nao existe: {model_path}")


def validate_strict_gate(output_path: Path, input_path: Path | None) -> None:
    if not output_requires_strict_gate(output_path):
        return

    if input_path is None:
        raise ValueError(
            "Saida em 'saida_docx' exige gate de fluxo estrito. "
            "Use arquivo-fonte UTF-8 e crie o sidecar '.gate.json' correspondente."
        )

    gate_path = derive_gate_path(input_path)
    if not gate_path.exists():
        raise ValueError(
            "Gate de fluxo estrito ausente para geracao do .docx. "
            f"Crie o arquivo: {gate_path}"
        )

    try:
        gate = json.loads(gate_path.read_text(encoding="utf-8-sig"))
    except json.JSONDecodeError as exc:
        raise ValueError(f"Gate de fluxo estrito invalido: {gate_path} ({exc})") from exc

    if not isinstance(gate, dict):
        raise ValueError(f"Gate de fluxo estrito deve ser um objeto JSON: {gate_path}")

    for key in STRICT_GATE_KEYS:
        if key == "peticoes_posteriores":
            require_non_empty_peticoes(gate)
        else:
            require_non_empty_string(gate, key)

    if gate.get("preflight_confirmado") is not True:
        raise ValueError("Gate de fluxo estrito sem confirmacao formal de preflight.")
    if gate.get("postflight_confirmado") is not True:
        raise ValueError("Gate de fluxo estrito sem confirmacao formal de postflight.")

    validate_model_base_path(require_non_empty_string(gate, "modelo_base_path"))


def add_paragraph(document: Document, raw_text: str) -> None:
    text = raw_text.rstrip("\n")
    if not text:
        document.add_paragraph("")
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(28)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15

    cursor = 0
    while True:
        start = text.find("**", cursor)
        if start == -1:
            run = paragraph.add_run(text[cursor:])
            set_run_font(run, "Times New Roman", 12)
            break
        end = text.find("**", start + 2)
        if end == -1:
            run = paragraph.add_run(text[cursor:])
            set_run_font(run, "Times New Roman", 12)
            break
        if start > cursor:
            run = paragraph.add_run(text[cursor:start])
            set_run_font(run, "Times New Roman", 12)
        bold_text = text[start + 2 : end]
        run = paragraph.add_run(bold_text)
        set_run_font(run, "Times New Roman", 12, bold=True)
        cursor = end + 2


def create_docx(output_path: Path, paragraphs: list[str], input_path: Path | None = None) -> None:
    validate_strict_gate(output_path, input_path)
    # validate_ptbr_text(paragraphs)
    document = Document()
    apply_normal_style(document)
    for paragraph in paragraphs:
        add_paragraph(document, paragraph)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    # saved = Document(output_path)
    # validate_ptbr_text([p.text for p in saved.paragraphs])


def main() -> int:
    if len(sys.argv) not in {2, 3}:
        print("Usage: make_docx.py OUTPUT_PATH [INPUT_FILE_UTF8]", file=sys.stderr)
        return 2

    output_path = Path(sys.argv[1])
    input_path: Path | None = None
    if len(sys.argv) == 3:
        input_path = Path(sys.argv[2])
        content = input_path.read_text(encoding="utf-8-sig")
    else:
        content = sys.stdin.read()
    normalized = content.replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = [chunk for chunk in re.split(r"\n{2,}", normalized) if chunk]
    create_docx(output_path, paragraphs, input_path=input_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
